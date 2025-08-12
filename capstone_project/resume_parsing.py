import os
from dotenv import load_dotenv
import streamlit as st
import PyPDF2
from docx import Document
import requests
from streamlit_lottie import st_lottie
import io
from google import genai
import re

# Load environment variables from .env file
load_dotenv()

API_KEY = os.getenv("GEMINI_API_KEY")
if not API_KEY:
    raise ValueError("GEMINI_API_KEY not found in environment variables. Please set it in your .env file.")

# Initialize Gemini client globally (it reads the API key from environment)
client = genai.Client()

# ----------- Gemini API call -----------

def gemini_call(prompt, model="gemini-2.5-flash", max_tokens=512):
    response = client.models.generate_content(
        model=model,
        contents=prompt
    )
    return response.text

# ----------- Tool functions -----------

def tailor_resume(resume, job_description):
    system_prompt = (
        "You are a professional resume optimizer. Tailor the resume to match the job description with relevant keywords, "
        "reordered skills, and ATS-friendly formatting."
    )
    prompt = f"{system_prompt}\n\nResume:\n{resume}\n\nJob Description:\n{job_description}\n\nProvide tailored resume text only."
    return gemini_call(prompt)

def score_ats(resume, job_description):
    system_prompt = (
        "You are an ATS scoring assistant. Given a tailored resume and job description, score the resume's ATS compatibility "
        "out of 100 and suggest improvements."
    )
    prompt = f"{system_prompt}\n\nResume:\n{resume}\n\nJob Description:\n{job_description}\n\nProvide a numeric score and improvement suggestions."
    return gemini_call(prompt)

def generate_cover_letter(resume, job_description):
    system_prompt = (
        "You are a professional cover letter writer. Given a tailored resume and job description, "
        "write a customized cover letter aligned with the resume and job."
    )
    prompt = f"{system_prompt}\n\nResume:\n{resume}\n\nJob Description:\n{job_description}\n\nWrite the cover letter."
    return gemini_call(prompt)

def verify_outputs(tailored_resume, ats_feedback, cover_letter, job_description):
    verification_prompt = (
        "You are a verification assistant checking if all of the following are done correctly:\n\n"
        "Tailored Resume:\n{tailored_resume}\n\n"
        "ATS Score & Feedback:\n{ats_feedback}\n\n"
        "Cover Letter:\n{cover_letter}\n\n"
        "Job Description:\n{job_description}\n\n"
        "Check if the tailored resume matches the job description well, the ATS score is reasonable, and the cover letter aligns with the resume and job description. "
        "If anything is missing or needs redoing, explain exactly what and suggest redoing those steps. Otherwise, respond with the single word 'ALL GOOD'."
    ).format(
        tailored_resume=tailored_resume,
        ats_feedback=ats_feedback,
        cover_letter=cover_letter,
        job_description=job_description,
    )
    return gemini_call(verification_prompt)

# ----------- Helpers ------------

def extract_text_pdf(file):
    reader = PyPDF2.PdfReader(file)
    text = ""
    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            text += page_text + "\n"
    return text.strip()

def extract_text_docx(file):
    doc = Document(file)
    return "\n".join(para.text for para in doc.paragraphs).strip()

def load_lottie(url):
    r = requests.get(url)
    if r.status_code != 200:
        return None
    return r.json()

def clean_resume_text(text):
    """
    Clean resume text by:
    - Removing empty lines
    - Removing lines with typical code or comments (e.g., starts with #, //, or contains code-like patterns)
    - Removing excessive whitespace
    """
    cleaned_lines = []
    for line in text.splitlines():
        line = line.strip()
        if not line:
            continue  # skip empty lines
        # Remove lines that look like code or comments
        if re.match(r"^(#|//|--|/\*|\*|import |from |def |class |return |print\()", line):
            continue
        # Optionally, remove lines with many special chars (basic heuristic)
        if len(re.findall(r"[{}<>;=]", line)) > 3:
            continue
        cleaned_lines.append(line)
    return "\n".join(cleaned_lines)
def text_to_bytes(text, ext="txt"):
    if ext == "txt":
        return io.BytesIO(text.encode('utf-8'))
    elif ext == "docx":
        from docx import Document as DocxDocument
        doc = DocxDocument()
        for line in text.split("\n"):
            doc.add_paragraph(line)
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        return file_stream
    elif ext == "pdf":
        try:
            from fpdf import FPDF
        except ImportError:
            return None
        pdf = FPDF()
        pdf.add_page()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.set_font("Arial", size=12)
        for line in text.split("\n"):
            # Use multi_cell instead of cell for better line wrapping
            pdf.multi_cell(0, 10, line.encode('latin-1', 'replace').decode('latin-1'))
        pdf_bytes = pdf.output(dest='S').encode('latin-1')  # Return PDF as string then encode
        return io.BytesIO(pdf_bytes)
    else:
        return io.BytesIO(text.encode('utf-8'))


# ----------- Main Streamlit UI ------------
def main():
    st.set_page_config(page_title="AI Resume Tailoring & Cover Letter Generator", layout="wide")
    st.title("🤖 AI Resume Tailoring & Cover Letter Generator")

    if 'generated' not in st.session_state:
        st.session_state.generated = False
    if 'tailored_resume' not in st.session_state:
        st.session_state.tailored_resume = ""
    if 'ats_score_feedback' not in st.session_state:
        st.session_state.ats_score_feedback = ""
    if 'cover_letter' not in st.session_state:
        st.session_state.cover_letter = ""

    uploaded_file = st.file_uploader("📄 Upload your Resume (PDF or DOCX)", type=["pdf", "docx"])

    if not st.session_state.generated:
        job_description = st.text_area("📝 Paste the Job Description here", height=200)
    else:
        st.markdown("### Job description input hidden after generation.")

    if uploaded_file is not None:
        try:
            if uploaded_file.type == "application/pdf":
                resume_text = extract_text_pdf(uploaded_file)
            else:
                resume_text = extract_text_docx(uploaded_file)
            resume_text = clean_resume_text(resume_text)
        except Exception as e:
            st.error(f"Failed to extract resume text: {e}")
            st.stop()
    else:
        resume_text = ""

    lottie_url = "https://assets4.lottiefiles.com/packages/lf20_j1adxtyb.json"
    lottie_json = load_lottie(lottie_url)

    if st.button("🚀 Generate Tailored Resume, ATS Score & Cover Letter") and not st.session_state.generated:
        if not resume_text or not job_description.strip():
            st.error("Please upload a resume and enter a job description.")
            st.stop()

        with st.spinner("🤖 Running AI pipeline..."):
            if lottie_json:
                st_lottie(lottie_json, speed=1, height=150, key="loading")

            tailored_resume = tailor_resume(resume_text, job_description)
            ats_score_feedback = score_ats(tailored_resume, job_description)
            cover_letter = generate_cover_letter(tailored_resume, job_description)
            verification_result = verify_outputs(tailored_resume, ats_score_feedback, cover_letter, job_description)

            retry_count = 0
            max_retries = 2
            while "ALL GOOD" not in verification_result.upper() and retry_count < max_retries:
                st.info(f"Verification failed, retrying AI outputs (attempt {retry_count + 1})...")
                tailored_resume = tailor_resume(resume_text, job_description)
                ats_score_feedback = score_ats(tailored_resume, job_description)
                cover_letter = generate_cover_letter(tailored_resume, job_description)
                verification_result = verify_outputs(tailored_resume, ats_score_feedback, cover_letter, job_description)
                retry_count += 1

            st.session_state.tailored_resume = tailored_resume
            st.session_state.ats_score_feedback = ats_score_feedback
            st.session_state.cover_letter = cover_letter
            st.session_state.generated = True

    if st.session_state.generated:
        st.markdown("## ✨ Tailored Resume")
        st.text_area("Tailored Resume", value=st.session_state.tailored_resume, height=300)

        st.markdown("## 📊 ATS Score & Feedback")
        st.text_area("ATS Score & Feedback", value=st.session_state.ats_score_feedback, height=200)

        st.markdown("## 📝 Cover Letter")
        st.text_area("Cover Letter", value=st.session_state.cover_letter, height=300)

        st.markdown("## 📥 Download your files")

        resume_bytes = text_to_bytes(st.session_state.tailored_resume, ext="pdf")
        cover_bytes = text_to_bytes(st.session_state.cover_letter, ext="pdf")

        st.download_button("Download Tailored Resume (PDF)", data=resume_bytes, file_name="tailored_resume.pdf", mime="application/pdf")
        st.download_button("Download Cover Letter (PDF)", data=cover_bytes, file_name="cover_letter.pdf", mime="application/pdf")

        resume_docx = text_to_bytes(st.session_state.tailored_resume, ext="docx")
        cover_docx = text_to_bytes(st.session_state.cover_letter, ext="docx")

        if resume_docx:
            st.download_button("Download Tailored Resume (DOCX)", data=resume_docx, file_name="tailored_resume.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

        if cover_docx:
            st.download_button("Download Cover Letter (DOCX)", data=cover_docx, file_name="cover_letter.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

if __name__ == "__main__":
    main()